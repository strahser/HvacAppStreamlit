import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_BREAK
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd


class DocxConfig:
	def __init__(self):
		"""
		create/open file/. Add Margins,boarders,Title,tables to Docx
		"""
		self.document = None

	def open_file(self, docx_file_path: str):
		doc_ = open(docx_file_path, "rb")
		self.document = Document(doc_)
		return self.document

	def create_file(self):
		self.document = Document()
		return self.document

	def add_page_borders(self):
		border_list =(
				"top",
				"left",
				"bottom",
				"right",
		)
		sec_pr = self.document.sections[0]._sectPr  # get the section properties el
		# create new borders el
		pg_borders = OxmlElement("w:pgBorders")
		# specifies how the relative positioning of the borders should be calculated
		pg_borders.set(qn("w:offsetFrom"), "page")
		for border_name in border_list:  # set all borders
			border_el = OxmlElement(f"w:{border_name}")
			border_el.set(qn("w:val"), "single")  # a single line
			border_el.set(
				qn("w:sz"), "4"
			)  # for meaning of  remaining attrs please look docs
			border_el.set(qn("w:space"), "24")
			border_el.set(qn("w:color"), "auto")
			pg_borders.append(border_el)  # register single border to border el
		sec_pr.append(pg_borders)  # apply border changes to section

	def add_margins(self, _height=297, _width=420):
		section = self.document.sections[0]
		section.page_height = Mm(_height)
		section.page_width = Mm(_width)
		section.left_margin = Mm(20)
		section.right_margin = Mm(5)
		section.top_margin = Mm(5)
		section.bottom_margin = Mm(5)
		section.header_distance = Mm(5)
		section.footer_distance = Mm(5)

	def __add_title_config(self, tabletitle):
		run = self.document.add_paragraph().add_run(tabletitle)
		font = run.font
		font.size = Pt(16)
		font.bold = True

	def add_doctable(self, data: pd.DataFrame, tabletitle: str):
		self.__add_title_config(tabletitle)

		table = self.document.add_table(
			data.shape[0] + 1, data.shape[1], style="Table Grid"
		)  # First row are table headers!

		table.allow_autofit = True
		table.autofit = True
		# add the header rows.
		for j in range(data.shape[-1]):
			if isinstance(data.columns[j], tuple):
				table.cell(0, j).text = data.columns[j][1]
				table.cell(1, j).text = "test"
			# add the bold font.
			else:
				table.cell(0, j).text = data.columns[j]

			table.cell(0, j).paragraphs[0].runs[0].font.bold = True
		# add the rest of the data frame
		for row in range(data.shape[0]):
			for col in range(data.shape[-1]):
				try:
					table.cell(row + 1, col).text = str(round(data.values[row, col], 2))
				except:
					table.cell(row + 1, col).text = str(data.values[row, col])

	def add_picture_table(self, first_row_list, picture_path, tabletitle):
		self.__add_title_config(tabletitle)

		table = self.document.add_table(rows=3, cols=len(first_row_list))

		for en, eq in enumerate(first_row_list):
			# --- by default a cell has one paragraph with zero runs ---
			table.cell(0, en).text = eq
			row = table.rows[1]
			row.cells[en].text = str(en)
			row.cells[en].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
			row.cells[en].paragraphs[0].runs[0].font.bold = True
			paragraph = table.cell(2, en).paragraphs[0]
			# --- add a run in which to place the picture ---
			run = paragraph.add_run()
			# --- add the picture to that run ---
			paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
			run.add_picture(picture_path[en], width=Mm(15), height=Mm(25))

	def add_page_break(self):
		self.document.add_page_break()

	def change_footer_text(self, new_text, footer_number=1):
		footer = self.document.sections[0].footer
		footer.paragraphs[footer_number].text = new_text

	def save(self, pathfile: str = "new"):
		self.document.save(pathfile)


class RenderDocx:

	def __init__(self, InputTableNames) -> None:
		self.doc_file = DocxConfig()
		self.InputTableNames = InputTableNames

	def _create_docx_file(self):
		self.doc_file.create_file()
		self.doc_file.add_margins(420, 297)

	def open_docx_file(self, docx_file_path: str):
		if os.path.isfile(docx_file_path):
			return self.doc_file.open_file(docx_file_path)

	def change_footer_text(self, new_text='new_text,footer_number=1', footer_number=1):
		self.doc_file.change_footer_text(new_text, footer_number)

	def create_or_open_docx_file(self, docx_file_path):
		if self.open_docx_file(docx_file_path):
			return self.open_docx_file(docx_file_path)
		else:
			return self._create_docx_file()

	def __add_pictures_ahu_table(self, context_dictionary):
		self.doc_file.add_picture_table(
			context_dictionary.ahu_equip_name,
			context_dictionary.ahu_pictures,
			self.InputTableNames.filter_pivot_table,
		)

	def add_context_to_file(self, context_dictionary):
		self.doc_file.document.add_heading(
			f'{self.InputTableNames.heading} {context_dictionary.system_name}',
			level=1)
		self.__add_pictures_ahu_table(context_dictionary)

		self.doc_file.add_doctable(
			context_dictionary.ahu_property,
			self.InputTableNames.ahu_equipment_property,
		)
		self.doc_file.add_doctable(
			context_dictionary.ahu_excel_df,
			self.InputTableNames.ahu_equipment_table_name,
		)

		self.doc_file.add_page_break()

	def change_footer_text(self,
	                       new_text='new_text,footer_number=1',
	                       footer_number=1):
		self.doc_file.change_footer_text(new_text, footer_number)

	def save(self, pathfile):
		self.doc_file.save(pathfile)
